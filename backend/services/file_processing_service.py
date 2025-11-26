import logging
import io
import base64
from typing import List, Dict, Optional, Tuple
from pathlib import Path
import tempfile
import httpx

logger = logging.getLogger(__name__)

class FileProcessingService:
    def __init__(self, hf_api_token: str = None, hf_base_url: str = None, hf_model: str = None):
        self.ocr_reader = None
        self.hf_api_token = hf_api_token
        self.hf_base_url = hf_base_url or "https://router.huggingface.co/v1/chat/completions"
        self.hf_model = hf_model or "deepseek-ai/DeepSeek-V3.2-Exp"

        # Supported file types
        self.supported_image_types = ['image/jpeg', 'image/jpg', 'image/png', 'image/gif', 'image/webp']
        self.supported_pdf_type = 'application/pdf'
        self.supported_excel_types = [
            'application/vnd.ms-excel',  # .xls
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'  # .xlsx
        ]
        self.supported_word_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
            'application/msword'  # .doc
        ]
        self.supported_text_types = ['text/plain', 'text/csv', 'application/csv']

    def _init_ocr(self):
        """Lazy initialization of OCR reader"""
        if self.ocr_reader is None:
            try:
                import easyocr
                # Support Thai and English
                self.ocr_reader = easyocr.Reader(['th', 'en'], gpu=False)
                logger.info("✅ OCR reader initialized (Thai + English)")
            except Exception as e:
                logger.error(f"Failed to initialize OCR: {e}")
                raise
        return self.ocr_reader

    async def process_files(self, files: List[Tuple[str, bytes, str]]) -> Dict[str, any]:
        """
        Process uploaded files (images, PDFs, Excel, Word, CSV, text)

        Args:
            files: List of (filename, content, content_type) tuples

        Returns:
            Dict with extracted text, images info, and file summaries
        """
        result = {
            'extracted_text': [],
            'images': [],
            'pdfs': [],
            'excel_files': [],
            'word_files': [],
            'text_files': [],
            'total_files': len(files),
            'success': True,
            'errors': []
        }

        for filename, content, content_type in files:
            try:
                if content_type in self.supported_image_types:
                    text = await self._process_image_deepseek(content, filename)
                    result['extracted_text'].append({
                        'source': filename,
                        'type': 'image',
                        'text': text
                    })
                    result['images'].append({
                        'filename': filename,
                        'size': len(content),
                        'has_text': bool(text)
                    })

                elif content_type == self.supported_pdf_type:
                    text = await self._process_pdf(content, filename)
                    result['extracted_text'].append({
                        'source': filename,
                        'type': 'pdf',
                        'text': text
                    })
                    result['pdfs'].append({
                        'filename': filename,
                        'size': len(content),
                        'pages': text.count('\n--- Page')
                    })

                elif content_type in self.supported_excel_types:
                    text = await self._process_excel(content, filename)
                    result['extracted_text'].append({
                        'source': filename,
                        'type': 'excel',
                        'text': text
                    })
                    result['excel_files'].append({
                        'filename': filename,
                        'size': len(content)
                    })

                elif content_type in self.supported_word_types:
                    text = await self._process_word(content, filename)
                    result['extracted_text'].append({
                        'source': filename,
                        'type': 'word',
                        'text': text
                    })
                    result['word_files'].append({
                        'filename': filename,
                        'size': len(content)
                    })

                elif content_type in self.supported_text_types or filename.endswith(('.txt', '.csv')):
                    text = await self._process_text(content, filename)
                    result['extracted_text'].append({
                        'source': filename,
                        'type': 'text',
                        'text': text
                    })
                    result['text_files'].append({
                        'filename': filename,
                        'size': len(content)
                    })

                else:
                    result['errors'].append(f"Unsupported file type: {filename} ({content_type})")

            except Exception as e:
                error_msg = f"Error processing {filename}: {str(e)}"
                logger.error(error_msg)
                result['errors'].append(error_msg)
                result['success'] = False

        return result

    async def _process_image_deepseek(self, image_bytes: bytes, filename: str) -> str:
        """Extract text from image using DeepSeek-V3.2-Exp (Vision Model)"""
        if not self.hf_api_token:
            logger.warning("No HF API token, falling back to easyOCR")
            return await self._process_image(image_bytes, filename)

        try:
            # Convert image to base64
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')

            # Determine image format from filename
            ext = Path(filename).suffix.lower().replace('.', '')
            if ext not in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
                ext = 'png'

            data_uri = f"data:image/{ext};base64,{image_base64}"

            payload = {
                "model": self.hf_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": data_uri}
                            },
                            {
                                "type": "text",
                                "text": "Extract all text from this image. Preserve the original language (Thai, English, Japanese, etc.). Return only the extracted text without any additional explanation."
                            }
                        ]
                    }
                ],
                "max_tokens": 2048,
                "temperature": 0.1
            }

            headers = {
                "Authorization": f"Bearer {self.hf_api_token}",
                "Content-Type": "application/json"
            }

            async with httpx.AsyncClient(timeout=60) as client:
                response = await client.post(
                    self.hf_base_url,
                    headers=headers,
                    json=payload
                )

                if response.status_code == 200:
                    data = response.json()
                    extracted_text = data["choices"][0]["message"]["content"].strip()
                    logger.info(f"DeepSeek-V3.2-Exp extracted text from {filename} ({len(extracted_text)} chars)")
                    return extracted_text
                else:
                    logger.error(f"DeepSeek-V3.2-Exp failed with status {response.status_code}: {response.text[:200]}")
                    # Fallback to easyOCR
                    return await self._process_image(image_bytes, filename)

        except Exception as e:
            logger.error(f"DeepSeek-V3.2-Exp failed for {filename}: {e}")
            # Fallback to easyOCR
            return await self._process_image(image_bytes, filename)

    async def _process_image(self, image_bytes: bytes, filename: str) -> str:
        """Extract text from image using EasyOCR (fallback)"""
        try:
            from PIL import Image

            # Load image
            image = Image.open(io.BytesIO(image_bytes))

            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Initialize OCR
            reader = self._init_ocr()

            # Perform OCR
            result = reader.readtext(image, detail=0)

            if result:
                extracted_text = '\n'.join(result)
                logger.info(f"EasyOCR extracted {len(result)} text blocks from {filename}")
                return extracted_text
            else:
                logger.warning(f"No text found in {filename}")
                return ""

        except Exception as e:
            logger.error(f"OCR failed for {filename}: {e}")
            raise

    async def _process_pdf(self, pdf_bytes: bytes, filename: str) -> str:
        """Extract text from PDF"""
        try:
            import pdfplumber

            extracted_text = []

            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        extracted_text.append(f"--- Page {page_num} ---\n{text}")

            if extracted_text:
                result = '\n\n'.join(extracted_text)
                logger.info(f"Extracted text from {len(pdf.pages)} pages in {filename}")
                return result
            else:
                logger.warning(f"No text found in PDF {filename}")
                return ""

        except Exception as e:
            logger.error(f"PDF processing failed for {filename}: {e}")
            raise

    async def _process_excel(self, excel_bytes: bytes, filename: str) -> str:
        """Extract text from Excel file (.xlsx, .xls)"""
        try:
            import pandas as pd

            extracted_text = []

            # Read all sheets
            excel_file = pd.ExcelFile(io.BytesIO(excel_bytes))

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                # Convert DataFrame to readable text
                sheet_text = f"--- Sheet: {sheet_name} ---\n"
                sheet_text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
                sheet_text += df.to_string(index=False)

                extracted_text.append(sheet_text)

            if extracted_text:
                result = '\n\n'.join(extracted_text)
                logger.info(f"Extracted {len(excel_file.sheet_names)} sheets from {filename}")
                return result
            else:
                logger.warning(f"No data found in Excel file {filename}")
                return ""

        except Exception as e:
            logger.error(f"Excel processing failed for {filename}: {e}")
            raise

    async def _process_word(self, word_bytes: bytes, filename: str) -> str:
        """Extract text from Word document (.docx)"""
        try:
            import docx

            doc = docx.Document(io.BytesIO(word_bytes))

            extracted_text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text.append(para.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables, start=1):
                table_text = f"\n--- Table {table_idx} ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                extracted_text.append(table_text)

            if extracted_text:
                result = '\n'.join(extracted_text)
                logger.info(f"Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables from {filename}")
                return result
            else:
                logger.warning(f"No text found in Word document {filename}")
                return ""

        except Exception as e:
            logger.error(f"Word processing failed for {filename}: {e}")
            raise

    async def _process_text(self, text_bytes: bytes, filename: str) -> str:
        """Extract text from plain text or CSV file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'tis-620']

            for encoding in encodings:
                try:
                    text = text_bytes.decode(encoding)

                    # Check if it's CSV
                    if filename.endswith('.csv'):
                        import csv
                        lines = text.splitlines()
                        reader = csv.reader(lines)
                        rows = list(reader)

                        if rows:
                            result = "--- CSV Data ---\n"
                            result += f"Rows: {len(rows)}, Columns: {len(rows[0]) if rows else 0}\n\n"
                            for row in rows[:100]:  # Limit to 100 rows for preview
                                result += " | ".join(row) + "\n"

                            logger.info(f"Extracted {len(rows)} rows from CSV {filename}")
                            return result
                    else:
                        logger.info(f"Extracted text from {filename} using {encoding} encoding")
                        return text

                except UnicodeDecodeError:
                    continue

            # If all encodings fail
            logger.warning(f"Could not decode {filename} with any encoding")
            return ""

        except Exception as e:
            logger.error(f"Text processing failed for {filename}: {e}")
            raise

    def create_chart(self, data: Dict, chart_type: str = 'bar') -> Optional[str]:
        """
        Create chart from data using Plotly

        Args:
            data: Dict with 'labels' and 'values' keys
            chart_type: Type of chart ('bar', 'line', 'pie')

        Returns:
            Base64 encoded PNG image of the chart
        """
        try:
            import plotly.graph_objects as go
            import plotly.io as pio

            labels = data.get('labels', [])
            values = data.get('values', [])

            if not labels or not values:
                logger.warning("No data provided for chart")
                return None

            if chart_type == 'bar':
                fig = go.Figure(data=[go.Bar(x=labels, y=values)])
            elif chart_type == 'line':
                fig = go.Figure(data=[go.Scatter(x=labels, y=values, mode='lines+markers')])
            elif chart_type == 'pie':
                fig = go.Figure(data=[go.Pie(labels=labels, values=values)])
            else:
                logger.warning(f"Unsupported chart type: {chart_type}")
                return None

            fig.update_layout(
                title=data.get('title', 'Data Visualization'),
                template='plotly_white',
                width=800,
                height=500
            )

            # Convert to PNG and encode to base64
            img_bytes = pio.to_image(fig, format='png')
            img_base64 = base64.b64encode(img_bytes).decode('utf-8')

            logger.info(f"Created {chart_type} chart successfully")
            return f"data:image/png;base64,{img_base64}"

        except Exception as e:
            logger.error(f"Chart creation failed: {e}")
            return None

    def summarize_file_content(self, processing_result: Dict) -> str:
        """Create a summary of processed files for AI context"""
        summary_parts = []

        if processing_result.get('images'):
            summary_parts.append(f"📷 Images: {len(processing_result['images'])} files")

        if processing_result.get('pdfs'):
            summary_parts.append(f"📄 PDFs: {len(processing_result['pdfs'])} files")

        if processing_result.get('extracted_text'):
            text_blocks = []
            for item in processing_result['extracted_text']:
                if item['text']:
                    text_blocks.append(f"\n[{item['source']}]:\n{item['text']}")
            if text_blocks:
                summary_parts.append("\n📝 Extracted Text:" + ''.join(text_blocks))

        if processing_result.get('errors'):
            summary_parts.append(f"\n⚠️ Errors: {'; '.join(processing_result['errors'])}")

        return '\n'.join(summary_parts) if summary_parts else "No content extracted from files."
